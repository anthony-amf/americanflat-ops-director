# -*- coding: utf-8 -*-
# Operations 90-Day Plan — revised from meeting notes, on-brand AF Design System
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0x0F,0x0F,0x0F); GREY3 = RGBColor(0x66,0x66,0x66)
GREY2 = RGBColor(0xB3,0xB3,0xB3); WHITE = RGBColor(0xFF,0xFF,0xFF)
BLACK_HEX="0F0F0F"; GREY1_HEX="E6E6E6"; ZEBRA_HEX="F2F2F2"
FONT="Glacial Indifference"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = Inches(0.85); sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.85)
CW = Inches(8.5) - Inches(1.7)

st = doc.styles['Normal']
st.font.name = FONT; st.font.size = Pt(11); st.font.color.rgb = BLACK
st.paragraph_format.space_after = Pt(0); st.paragraph_format.line_spacing = 1.18

def shade(cell,hexfill):
    tcPr=cell._tc.get_or_add_tcPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),hexfill); tcPr.append(s)

def borders(cell,color=GREY1_HEX,sz="4"):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right'):
        x=OxmlElement('w:'+e); x.set(qn('w:val'),'single'); x.set(qn('w:sz'),sz)
        x.set(qn('w:space'),'0'); x.set(qn('w:color'),color); b.append(x)
    tcPr.append(b)

def margins(cell,t=80,b=80,l=120,r=120):
    tcPr=cell._tc.get_or_add_tcPr(); m=OxmlElement('w:tcMar')
    for k,v in (('top',t),('bottom',b),('left',l),('right',r)):
        n=OxmlElement('w:'+k); n.set(qn('w:w'),str(v)); n.set(qn('w:type'),'dxa'); m.append(n)
    tcPr.append(m)

def cellstyle(cell,fill=None,t=80,b=80,l=120,r=120,va=None,bcolor=GREY1_HEX):
    borders(cell,bcolor); margins(cell,t,b,l,r)
    if fill: shade(cell,fill)
    if va: cell.vertical_alignment=va

def run(p,text,size=11,color=BLACK,bold=False):
    r=p.add_run(text); r.font.name=FONT; r.font.size=Pt(size); r.font.color.rgb=color; r.font.bold=bold
    rpr=r._element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rpr.insert(0,rf)
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a),FONT)
    return r

def para(after=4,before=0,ls=1.18):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.line_spacing=ls
    return p

def cpara(cell,after=0,before=0,ls=1.12):
    p=cell.add_paragraph(); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.line_spacing=ls
    return p

def hairline(p,color=BLACK_HEX,sz="6",space="3",edge='bottom'):
    pPr=p._p.get_or_add_pPr(); bd=OxmlElement('w:pBdr'); e=OxmlElement('w:'+edge)
    e.set(qn('w:val'),'single'); e.set(qn('w:sz'),sz); e.set(qn('w:space'),space); e.set(qn('w:color'),color)
    bd.append(e); pPr.append(bd)

def colwidths(tbl,widths):
    tbl.autofit=False; tbl.allow_autofit=False
    for row in tbl.rows:
        for i,w in enumerate(widths): row.cells[i].width=w

def section_head(num,eyebrow,title):
    p=para(after=1,before=14); run(p,f"{num} — {eyebrow}",size=8.5,color=GREY3)
    p2=para(after=5,before=0); run(p2,title,size=15,color=BLACK,bold=True)
    hairline(p2,color=BLACK_HEX,sz="6")

def body(text,after=6):
    p=para(after=after); run(p,text,size=11,color=BLACK)

def ai_score(cell,score):
    p=cpara(cell,after=0,before=2,ls=1.0)
    run(p,"AI  ",size=8.5,color=GREY3); run(p,str(score),size=14,color=BLACK,bold=True); run(p," / 10",size=8.5,color=GREY3)
    bar=cpara(cell,after=0,before=1,ls=1.0)
    run(bar,"█"*score,size=7,color=BLACK); run(bar,"█"*(10-score),size=7,color=RGBColor(0xE6,0xE6,0xE6))

# footer
sec.different_first_page_header_footer = True
sec.first_page_header.paragraphs[0].text=""
hp2=sec.header.paragraphs[0]; hp2.alignment=WD_ALIGN_PARAGRAPH.LEFT
run(hp2,"americanflat",size=9,color=BLACK,bold=True)
for foot in (sec.footer.paragraphs[0], sec.first_page_footer.paragraphs[0]):
    foot.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run(foot,"Operations · 90-Day Plan · June 2026",size=8,color=GREY3)

# ============================================================
# PAGE 1
# ============================================================
p=para(after=2); run(p,"americanflat",size=13,color=BLACK,bold=True)
p=para(after=2,before=8); run(p,"Operations",size=27,color=BLACK,bold=True)
p=para(after=2); run(p,"A 90-Day Plan — Team AI Readiness & Development",size=14,color=GREY3)
p=para(after=4); run(p,"Prepared by Anthony Armstrong   ·   June 2026",size=9.5,color=GREY3)
hairline(p,color=BLACK_HEX,sz="8")

section_head("01","THE SITUATION","Where we are.")
body("Since January, the operations team has gone all-in on AI. Three team members are already building and shipping skills; three are still ramping — they understand the capability but haven't yet led an end-to-end AI workflow. The goal for the next 90 days is to move the entire team to ownership, where each person builds and operates AI agents in their domain, supported by a lean-AI mindset and clear contingency planning.")
body("This plan is a direct, honest read on where each person is today, what holds them back, and what we will deliver. AI score: 1–3 fully manual or single-tool use  ·  4–6 regular use for analysis  ·  7–8 building & shipping skills  ·  9–10 owns the platform and unblocks others.")

section_head("02","THE CONTEXT","The environment we are operating in.")
ctx=[("TEAM","Seven people: three builders (Anthony, Iván, John), three ramping (Carolina, Umer, Erickson, Nica)."),
     ("AI MATURITY","Builders drive the transformation. Ramping team members understand AI capabilities but lack hands-on ownership of an AI workflow."),
     ("SYSTEMS","Amazon VC, Luminous, STEDI/EDI, BigQuery pipelines, Claude agents, GitHub skill-publishing infra, Slack bots."),
     ("RISK","Single point of failure on John for WH operations and skill-building. Key-person concentration on Iván (platform) and Anthony (vendor recovery).")]
t=doc.add_table(rows=2,cols=2); half=Emu(int(CW)//2)
for i,(lab,bod) in enumerate(ctx):
    c=t.cell(i//2,i%2); cellstyle(c,fill=None,t=100,b=100,l=140,r=140)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.0
    run(p,lab,size=8.5,color=GREY3,bold=True)
    bp=cpara(c,after=0,ls=1.15); run(bp,bod,size=10,color=BLACK)
colwidths(t,[half,half])

section_head("03","THE PLAN","Three phases. Ninety days.")
phases=[("01","Days 1–30","Stabilize","Erickson and Nica join a real daily task (e.g. BTR skill) with John and Carolina. Document critical knowledge: EDI maps, data pipelines, VC workflows."),
        ("02","Days 31–60","Systematize","Each team member owns a Claude agent in their domain. Umer, Erickson, Nica ship their first end-to-end workflow. Light skill-review standard for quality gate."),
        ("03","Days 61–90","Scale","Build contingency for John (WH ops backup). Set human-in-the-loop controls for billing and acceptance (Luminous + Claude + human approval). Revisit headcount post Day 90.")]
t=doc.add_table(rows=1,cols=3); third=Emu(int(CW)//3)
for i,(num,days,name,bod) in enumerate(phases):
    c=t.cell(0,i); cellstyle(c,t=110,b=110,l=130,r=130)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1.0
    run(p,num,size=16,color=BLACK,bold=True); run(p,f"   {days}",size=8,color=GREY3)
    np_=cpara(c,after=3,ls=1.0); run(np_,name,size=12,color=BLACK,bold=True)
    bp=cpara(c,after=0,ls=1.15); run(bp,bod,size=9,color=BLACK)
colwidths(t,[third,third,third])

# ============================================================
# PAGE 2 — TEAM ASSESSMENT
# ============================================================
doc.add_page_break()
section_head("04","THE TEAM","Seven people. AI ratings and key gaps.")

PEOPLE=[
 dict(name="Anthony Armstrong", role="Director of Operations", real="Owns AI transformation, vendor recovery ($1.2M+), 3PL strategy", ai=10,
   m1="Owns the transformation; no one else has the bandwidth or mandate.", c1="Vast surface area: AI adoption, vendor contracts, IT admin, day-to-day ops.", p1="Key risk if he steps back. Needs explicit delegation and scope recognition."),
 dict(name="Iván Calderón", role="AI & Data Manager", real="Owns data platform, GCP/BigQuery, skill-publishing infra, identity / access", ai=10,
   m1="Built the entire data layer and skill-publishing system. No one understands the platform.", c1="Single point of failure across infra and security. Daily firefighting eats build time.", p1="Needs documented backup and protected build time. Platform-owner mandate to route requests through systems."),
 dict(name="John Nunez", role="Ops Manager", real="Owns WH flow, VC compliance, builds Claude skills, runs 3S-style improvement", ai=8,
   m1="Shipping end-to-end skills (Short Processor, VC-PO bot, Walmart companion). Strong practitioner.", c1="WH daily firefighting competes with builder time. No backup for critical WH ops.", p1="Likely reaches 10/10 after 30-day training. Key contingency risk — if John leaves, WH ops + AI building go dark."),
 dict(name="Carolina del Rio", role="People Ops & Systems Analyst", real="Really: EDI/Integrations Engineer, automation builder, people-ops systems", ai=7,
   m1="Built the Attendance app, 940 validator, Coffee scheduler, Luminous report bot. Giorgio calls her an expert.", c1="Still ramping on GitHub deploys (leans on Iván); EDI knowledge concentrated in her.", p1="Needs pairing with Iván for self-sufficient deploys. Good candidate to help onboard Erickson/Nica to building."),
 dict(name="Umer Raja", role="Ops Assistant — VC Compliance", real="Chargeback disputes (36–258/day), SIOC certification, orders reconciliation, Zendesk", ai=3,
   m1="Built an agent and summary tool; slightly above zero hands-on use.", c1="Narrow mindset; rigid schedule; doesn't think outside scope. Works in isolation.", p1="Needs the BTR shared-task model (daily, with others, learning by doing). High-ROI automation target; flagged as most-automatable role if headcount reduction needed."),
 dict(name="Erickson Dela Cruz", role="Ops Assistant — VC Acceptance", real="VC order acceptance (US/EU/CN), Andon Cord complaint process, price/shortage claims", ai=3,
   m1="Knows AI capabilities from meetings. Consistent, reliable, never misses an update.", c1="Zero hands-on AI use. Doesn't see how it applies to his repetitive VC-tracker work.", p1="Start with BTR task (daily shared model). Partner with Carolina/John. High-upside target: VC-tracker entry could be 80% automated."),
 dict(name="Nica Jordan", role="Ops Assistant — VC Invoicing", real="VC invoicing pipeline, Invoice Recap to leadership, QB/Luminous sync, invoice validation", ai=2,
   m1="Understands capabilities; reports billed revenue to the CEO — the stakes are real.", c1="Zero hands-on AI use yet. Narrow scope, not visible growth. Terse comms make blockers invisible.", p1="Needs the BTR shared task to build confidence. First win: invoice validation / FCR bot. Weekly checkpoints so progress is visible."),
]

C0,C1,C2=Twips(2700),Twips(3960),Twips(3236)
tbl=doc.add_table(rows=1,cols=3)
for cell,lab in zip(tbl.rows[0].cells,["PERSON & DOMAIN","MINDSET","CAPACITY / PERSONALITY"]):
    cellstyle(cell,fill=BLACK_HEX,t=60,b=60,l=120,r=120,va=WD_ALIGN_VERTICAL.CENTER,bcolor=BLACK_HEX)
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
    run(p,lab,size=8.5,color=WHITE,bold=True)

for i,pr in enumerate(PEOPLE):
    fill="FFFFFF" if i%2==0 else ZEBRA_HEX
    cells=tbl.add_row().cells
    # left cell: name, role, domain, AI
    c=cells[0]; cellstyle(c,fill=fill,t=80,b=80,l=120,r=110)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1.0
    run(p,pr["name"],size=10,color=BLACK,bold=True)
    p2=cpara(c,after=1,ls=1.05); run(p2,pr["role"],size=8,color=GREY3)
    p3=cpara(c,after=2,ls=1.05); run(p3,pr["real"],size=8,color=BLACK)
    ai_score(c,pr["ai"])
    # middle: mindset
    c1=cells[1]; cellstyle(c1,fill=fill,t=80,b=80,l=120,r=120)
    p=c1.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.12
    run(p,pr["m1"],size=8.5,color=BLACK)
    # right: capacity + personality
    c2=cells[2]; cellstyle(c2,fill=fill,t=80,b=80,l=120,r=120)
    p=c2.paragraphs[0]; p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.12
    run(p,"Capacity.  ",size=8.5,color=BLACK,bold=True); run(p,pr["c1"],size=8.5,color=BLACK)
    pn=cpara(c2,after=0,ls=1.12); run(pn,"Personality.  ",size=8.5,color=BLACK,bold=True); run(pn,pr["p1"],size=8.5,color=BLACK)

colwidths(tbl,[C0,C1,C2])

# ============================================================
# PAGE 3 — AI OWNERSHIP MODEL + MILESTONES + ASK
# ============================================================
doc.add_page_break()
section_head("05","AI OWNERSHIP MODEL","Future state. Each person owns agents in their domain.")
body("Automation: Each team member builds and operates Claude agents in their area — VC acceptance, invoicing, chargeback disputes, EDI, warehouse operations. Agents run daily, unattended, with human sign-off on exceptions.")
body("Human-in-the-loop controls: Billing (invoices) and acceptance (VC orders) stay under human review. Luminous + Claude do the heavy lifting; a person makes the final approval. Full automation creates unacceptable risk if that person is out.")
body("Contingency: The biggest risk is John. If he leaves, WH operations and skill-building both go dark. By Day 90, identify a backup for his WH responsibilities and document the skill-building process so someone else can take over if needed. Same for Iván's platform and Anthony's vendor relationships.")

section_head("06","MILESTONES","What we will deliver.")
MILE=[("Days 1–30","Erickson & Nica join BTR shared task with John & Carolina","John, Carolina","Daily participation, tasks completed together"),
      ("Days 1–30","EDI, VC, and pipeline workflows documented","Iván, Carolina","Written runbooks; a backup can execute them"),
      ("Days 31–60","Umer ships chargeback / shortage dispute bot","Umer, John","First end-to-end AI workflow in production"),
      ("Days 31–60","Erickson ships VC-tracker entry bot","Erickson, Carolina","AI-assisted acceptance workflow live"),
      ("Days 31–60","Nica ships invoice validation / FCR bot","Nica, Iván","Invoicing quality checks automated"),
      ("Days 61–90","John has a documented WH ops backup","John, Anthony","Cross-trained person can cover basic WH ops"),
      ("Days 61–90","Human-in-the-loop controls (Luminous + Claude) documented","Iván, John","Process for approval gates written, tested"),
      ("Days 61–90","Headcount and role review","Anthony","Revisit team structure post-training; decide on Umer role")]
cols=[Twips(1500),Twips(4000),Twips(1950),Twips(2286)]
mt=doc.add_table(rows=1,cols=4)
for cell,lab in zip(mt.rows[0].cells,["PHASE","DELIVERABLE","OWNER","MEASURE"]):
    cellstyle(cell,fill=BLACK_HEX,t=50,b=50,l=120,r=120,bcolor=BLACK_HEX)
    p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
    run(p,lab,size=8,color=WHITE,bold=True)
for i,m in enumerate(MILE):
    fill="FFFFFF" if i%2==0 else ZEBRA_HEX
    cells=mt.add_row().cells
    for j,v in enumerate(m):
        cellstyle(cells[j],fill=fill,t=55,b=55,l=120,r=120)
        p=cells[j].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.05
        run(p,v,size=9,color=BLACK,bold=(j==0))
colwidths(mt,cols)

section_head("07","THE ASK","What I am asking for.")
asks=[("1. Training Model","Let Erickson and Nica join the BTR shared task with John and Carolina, daily. Treat it like 3S — continuous improvement, everyone learning together, building real workflows, not classroom training."),
      ("2. Lean AI Mindset","Coach the team to think 'lean AI': start with a real problem, build the smallest workflow that solves it, iterate. Avoid over-engineering or waiting for 'perfect' automation."),
      ("3. Contingency Planning","By Day 90, document John's WH processes and skill-building approach so someone else can step in. Same for Iván's platform and Anthony's vendor relationships. The team is too concentrated.")]
for i,(num_and_title,bod) in enumerate(asks,1):
    p=para(after=5); run(p,num_and_title + ".  ",size=11,color=BLACK,bold=True)
    run(p,bod,size=11,color=BLACK)

p=para(after=2,before=10); hairline(p,color=BLACK_HEX,sz="8",edge='top',space="6")
run(p,"Give us 90 days.",size=14,color=BLACK,bold=True)
p=para(after=0,before=2)
run(p,"We will move the entire team from skill-builders to AI owners, each running agents in their domain. We will document the contingencies that keep the operation running if someone leaves. And we will establish a lean-AI culture where the team thinks in terms of small, iterative workflows, not big transformations.",size=11,color=BLACK)

out="/Users/aarmstrong/Claude-Amazon-AMF/Operations 90-Day Plan.docx"
doc.save(out); print("saved",out)
