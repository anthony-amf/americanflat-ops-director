# -*- coding: utf-8 -*-
# Operations 90-Day Plan — VISUAL, ADHD-friendly version
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0x0F,0x0F,0x0F); GREY3 = RGBColor(0x66,0x66,0x66)
GREY1 = RGBColor(0xE6,0xE6,0xE6); WHITE = RGBColor(0xFF,0xFF,0xFF)
BLACK_HEX="0F0F0F"; GREY1_HEX="E6E6E6"; LIGHT_HEX="F5F5F5"
FONT="Glacial Indifference"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = Inches(0.7); sec.bottom_margin = Inches(0.6)
sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)

st = doc.styles['Normal']
st.font.name = FONT; st.font.size = Pt(10); st.font.color.rgb = BLACK
st.paragraph_format.space_after = Pt(0); st.paragraph_format.line_spacing = 1.2

def shade(cell,hexfill):
    tcPr=cell._tc.get_or_add_tcPr(); s=OxmlElement('w:shd')
    s.set(qn('w:val'),'clear'); s.set(qn('w:color'),'auto'); s.set(qn('w:fill'),hexfill); tcPr.append(s)

def borders(cell,color=GREY1_HEX,sz="2"):
    tcPr=cell._tc.get_or_add_tcPr(); b=OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right'):
        x=OxmlElement('w:'+e); x.set(qn('w:val'),'single'); x.set(qn('w:sz'),sz)
        x.set(qn('w:space'),'0'); x.set(qn('w:color'),color); b.append(x)
    tcPr.append(b)

def margins(cell,t=120,b=120,l=140,r=140):
    tcPr=cell._tc.get_or_add_tcPr(); m=OxmlElement('w:tcMar')
    for k,v in (('top',t),('bottom',b),('left',l),('right',r)):
        n=OxmlElement('w:'+k); n.set(qn('w:w'),str(v)); n.set(qn('w:type'),'dxa'); m.append(n)
    tcPr.append(m)

def cellstyle(cell,fill=None,t=120,b=120,l=140,r=140,va=None,border_col=GREY1_HEX):
    borders(cell,border_col); margins(cell,t,b,l,r)
    if fill: shade(cell,fill)
    if va: cell.vertical_alignment=va

def run(p,text,size=10,color=BLACK,bold=False):
    r=p.add_run(text); r.font.name=FONT; r.font.size=Pt(size); r.font.color.rgb=color; r.font.bold=bold
    rpr=r._element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rpr.insert(0,rf)
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a),FONT)
    return r

def para(after=3,before=0,ls=1.15,align=None):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.line_spacing=ls
    if align: p.alignment=align
    return p

def cpara(cell,after=0,before=0,ls=1.15):
    p=cell.add_paragraph(); p.paragraph_format.space_after=Pt(after)
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.line_spacing=ls
    return p

def big_heading(text,size=32):
    p=para(after=6,before=12); run(p,text,size=size,color=BLACK,bold=True)
    return p

def subhead(text,size=16):
    p=para(after=4,before=8); run(p,text,size=size,color=BLACK,bold=True)
    return p

def callout_box(text,bg=BLACK_HEX,fg_color=WHITE):
    """Dark callout box with centered text"""
    t=doc.add_table(rows=1,cols=1); t.autofit=False; t.allow_autofit=False
    cell=t.rows[0].cells[0]
    cellstyle(cell,fill=bg,t=100,b=100,l=120,r=120,va=WD_ALIGN_VERTICAL.CENTER,border_col=BLACK_HEX)
    p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0)
    run(p,text,size=18,color=fg_color,bold=True)
    t.rows[0].cells[0].width=Emu(int(Inches(7.1)))
    para(after=8)

def ai_card(name,role,score,mindset_short,capacity_short):
    """Single team member visual card"""
    t=doc.add_table(rows=1,cols=2); t.autofit=False
    c0,c1=t.rows[0].cells

    # Left: name, role, score bar
    cellstyle(c0,fill=LIGHT_HEX,t=100,b=100,l=110,r=80,va=WD_ALIGN_VERTICAL.TOP)
    p=c0.paragraphs[0]; p.paragraph_format.space_after=Pt(1); p.paragraph_format.line_spacing=1.0
    run(p,name,size=11,color=BLACK,bold=True)
    p2=cpara(c0,after=4,ls=1.0); run(p2,role,size=8,color=GREY3)
    # AI bar
    p3=cpara(c0,after=0,ls=1.0); run(p3,"AI  ",size=8,color=GREY3); run(p3,str(score),size=16,color=BLACK,bold=True); run(p3,"/10",size=7,color=GREY3)
    barP=cpara(c0,after=0,ls=1.0); run(barP,"█"*score,size=6,color=BLACK); run(barP,"█"*(10-score),size=6,color=RGBColor(0xE0,0xE0,0xE0))

    # Right: mindset + capacity bullets
    cellstyle(c1,fill=LIGHT_HEX,t=100,b=100,l=80,r=110,va=WD_ALIGN_VERTICAL.TOP)
    mp=c1.paragraphs[0]; mp.paragraph_format.space_after=Pt(1); mp.paragraph_format.line_spacing=1.1
    run(mp,"Mindset: ",size=8,color=BLACK,bold=True); run(mp,mindset_short,size=8,color=BLACK)
    cp=cpara(c1,after=0,ls=1.1); run(cp,"Capacity: ",size=8,color=BLACK,bold=True); run(cp,capacity_short,size=8,color=BLACK)

    set_col_widths(t,[Twips(2400),Twips(4000)])
    para(after=2)

def set_col_widths(tbl,widths):
    tbl.autofit=False; tbl.allow_autofit=False
    for row in tbl.rows:
        for i,w in enumerate(widths): row.cells[i].width=w

# FOOTER
sec.different_first_page_header_footer=True
sec.first_page_header.paragraphs[0].text=""
hp=sec.header.paragraphs[0]; hp.alignment=WD_ALIGN_PARAGRAPH.LEFT
run(hp,"americanflat",size=8,color=BLACK,bold=True)
for foot in (sec.footer.paragraphs[0], sec.first_page_footer.paragraphs[0]):
    foot.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run(foot,"Operations · 90-Day Plan · June 2026",size=7,color=GREY3)

# ============================================================
# PAGE 1 — VISUAL OVERVIEW
# ============================================================
p=para(after=1); run(p,"americanflat",size=11,color=BLACK,bold=True)
big_heading("Operations", 28)
p=para(after=2); run(p,"90-Day AI Readiness Plan",size=13,color=GREY3)
p=para(after=8); run(p,"June 2026",size=9,color=GREY3)

# 7 people, quick snapshot
subhead("The Team Right Now")
snapshot_tbl=doc.add_table(rows=2,cols=4); snapshot_tbl.autofit=False
names_scores=[("Anthony","10"),("Iván","10"),("John","8"),("Carolina","7")]
for i,(nm,sc) in enumerate(names_scores):
    c=snapshot_tbl.cell(0,i); cellstyle(c,fill=BLACK_HEX,t=60,b=60,l=60,r=60,va=WD_ALIGN_VERTICAL.CENTER,border_col=BLACK_HEX)
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0)
    run(p,nm,size=8,color=WHITE,bold=True)
    c2=snapshot_tbl.cell(1,i); cellstyle(c2,fill=LIGHT_HEX,t=80,b=80,l=60,r=60,va=WD_ALIGN_VERTICAL.CENTER)
    p2=c2.paragraphs[0]; p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after=Pt(0)
    run(p2,sc,size=18,color=BLACK,bold=True)
set_col_widths(snapshot_tbl,[Twips(1600),Twips(1600),Twips(1600),Twips(1600)])

snapshot_tbl2=doc.add_table(rows=2,cols=3); snapshot_tbl2.autofit=False
names_scores2=[("Umer","3"),("Erickson","3"),("Nica","2")]
for i,(nm,sc) in enumerate(names_scores2):
    c=snapshot_tbl2.cell(0,i); cellstyle(c,fill=BLACK_HEX,t=60,b=60,l=60,r=60,va=WD_ALIGN_VERTICAL.CENTER,border_col=BLACK_HEX)
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0)
    run(p,nm,size=8,color=WHITE,bold=True)
    c2=snapshot_tbl2.cell(1,i); cellstyle(c2,fill=LIGHT_HEX,t=80,b=80,l=60,r=60,va=WD_ALIGN_VERTICAL.CENTER)
    p2=c2.paragraphs[0]; p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after=Pt(0)
    run(p2,sc,size=18,color=BLACK,bold=True)
set_col_widths(snapshot_tbl2,[Twips(2000),Twips(2000),Twips(2000)])
para(after=10)

# The problem
subhead("The Gap")
callout_box("3 builders  |  3 ramping  |  1 director", BLACK_HEX)
p=para(after=1); run(p,"Builders ship skills. Ramping team: understands AI but hasn't led a workflow yet.",size=10,color=BLACK)
para(after=8)

# Three phases visual
subhead("90 Days. Three Moves.")
phases_data=[
    ("Days\n1–30","Stabilize","Erickson & Nica join BTR task daily with John & Carolina.\nDocument critical knowledge."),
    ("Days\n31–60","Systematize","Each person owns an AI agent in their domain.\nShip first end-to-end workflows."),
    ("Days\n61–90","Scale","Contingency planning for John (WH backup).\nHeadcount & role review.")
]
phases_tbl=doc.add_table(rows=1,cols=3); phases_tbl.autofit=False
for i,(days,phase,desc) in enumerate(phases_data):
    c=phases_tbl.cell(0,i)
    cellstyle(c,fill=LIGHT_HEX,t=100,b=100,l=100,r=100,va=WD_ALIGN_VERTICAL.CENTER)
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.1
    run(p,days,size=9,color=GREY3); cp=cpara(c,after=4,ls=1.1); run(cp,phase,size=14,color=BLACK,bold=True)
    dp=cpara(c,after=0,ls=1.1); run(dp,desc,size=8,color=BLACK)
set_col_widths(phases_tbl,[Twips(2000),Twips(2000),Twips(2000)])
para(after=10)

# Key risks
subhead("Key Risks")
risks=[("John","WH ops + skill-building both go dark if he leaves."),
       ("Iván","Platform + security in one head. Daily firefighting eats build time."),
       ("Anthony","Vendor recovery + IT admin + AI adoption. Vast surface area.")]
for risk_person,risk_desc in risks:
    t=doc.add_table(rows=1,cols=2); t.autofit=False
    c0,c1=t.rows[0].cells
    cellstyle(c0,fill=BLACK_HEX,t=80,b=80,l=90,r=60); cellstyle(c1,fill=LIGHT_HEX,t=80,b=80,l=60,r=100)
    p0=c0.paragraphs[0]; p0.paragraph_format.space_after=Pt(0)
    run(p0,risk_person,size=10,color=WHITE,bold=True)
    p1=c1.paragraphs[0]; p1.paragraph_format.space_after=Pt(0); p1.paragraph_format.line_spacing=1.15
    run(p1,risk_desc,size=9,color=BLACK)
    set_col_widths(t,[Twips(1200),Twips(4700)])
para(after=8)

# ============================================================
# PAGE 2 — TEAM DETAIL
# ============================================================
doc.add_page_break()
big_heading("Team Cards", 22)
para(after=8)

team_data=[
    ("Anthony Armstrong","Director",10,"Owns the transformation mandate.","Vast surface area. Needs delegation."),
    ("Iván Calderón","Platform Lead",10,"Built the entire data layer.","Single point of failure. Firefighting eats build time."),
    ("John Nunez","Ops Manager",8,"Shipping end-to-end skills.","WH firefighting competes with builder time. KEY CONTINGENCY RISK."),
    ("Carolina del Rio","EDI Engineer",7,"Giorgio calls her an expert. Builds bots.","EDI knowledge concentrated with her. Ramping on GitHub deploys."),
    ("Umer Raja","Compliance",3,"Built an agent & summary tool.","Narrow mindset. Works in isolation. Highest automation upside."),
    ("Erickson Dela Cruz","Acceptance",3,"Knows AI capabilities. Consistent, reliable.","Zero hands-on AI yet. Doesn't see how it applies to his work."),
    ("Nica Jordan","Invoicing",2,"Understands capabilities. Stakes are real.","Zero hands-on AI. Narrow scope. Not visible growth."),
]

for name,role,score,mindset,capacity in team_data:
    ai_card(name,role,score,mindset,capacity)

# ============================================================
# PAGE 3 — THE ASK + MILESTONES
# ============================================================
doc.add_page_break()
big_heading("The Ask", 22)
para(after=6)

asks=[
    ("1. Training Model","Erickson & Nica join BTR with John & Carolina, daily.\nTreat it like 3S — continuous improvement, learning together."),
    ("2. Lean AI Mindset","Start with a real problem. Build the smallest workflow that solves it. Iterate.\nAvoid over-engineering or waiting for 'perfect' automation."),
    ("3. Contingency Planning","By Day 90, document John's WH processes so someone else can step in.\nSame for Iván's platform and Anthony's vendor relationships."),
]

for num_title,desc in asks:
    t=doc.add_table(rows=1,cols=2); t.autofit=False
    c0,c1=t.rows[0].cells
    cellstyle(c0,fill=BLACK_HEX,t=90,b=90,l=100,r=70); cellstyle(c1,fill=LIGHT_HEX,t=90,b=90,l=70,r=110)
    p0=c0.paragraphs[0]; p0.paragraph_format.space_after=Pt(0); p0.paragraph_format.line_spacing=1.1
    run(p0,num_title,size=10,color=WHITE,bold=True)
    p1=c1.paragraphs[0]; p1.paragraph_format.space_after=Pt(0); p1.paragraph_format.line_spacing=1.15
    run(p1,desc,size=9,color=BLACK)
    set_col_widths(t,[Twips(1400),Twips(4500)])
para(after=10)

# Key milestones (visual checklist)
subhead("Key Milestones")
milestones=[
    "☐  Erickson & Nica on BTR daily (Days 1–30)",
    "☐  EDI, VC, pipeline runbooks done (Days 1–30)",
    "☐  Umer ships chargeback bot (Days 31–60)",
    "☐  Erickson ships VC-tracker bot (Days 31–60)",
    "☐  Nica ships invoice validation bot (Days 31–60)",
    "☐  John has a documented WH backup (Days 61–90)",
    "☐  Human-in-loop controls documented (Days 61–90)",
    "☐  Headcount & role review (Days 61–90)",
]

for m in milestones:
    p=para(after=2); run(p,m,size=10,color=BLACK)

para(after=12)

# Big close
callout_box("Give us 90 days.", BLACK_HEX)
p=para(after=0); run(p,"We will move the entire team from skill-builders to AI owners, each running agents in their domain. We will document the contingencies that keep the operation running if someone leaves. And we will establish a lean-AI culture.",size=10,color=BLACK)

out="/Users/aarmstrong/Claude-Amazon-AMF/Operations 90-Day Plan — Visual.docx"
doc.save(out); print("saved",out)
